#!/usr/bin/env python3
"""
Script simple pour créer un template Word basique
"""

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    print("✅ Module python-docx disponible")
except ImportError:
    print("❌ Module python-docx non disponible")
    print("Installation: pip install python-docx")
    exit(1)

def create_simple_template():
    """Créer un template Word simple"""
    
    # Créer un nouveau document
    doc = Document()
    
    # Titre
    title = doc.add_heading('FORMULARIO DE PEDIDO DE PASSAPORTE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Espace
    doc.add_paragraph()
    
    # Section données personnelles
    doc.add_heading('Dados pessoais:', level=1)
    
    # Tableau simple
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    # Données
    data = [
        ('Nome:', '{{Prénom}}'),
        ('Apelido:', '{{Nom de famille}}'),
        ('Data de nascimento:', '{{Date de naissance}}'),
        ('Local de nascimento:', '{{Local de nascimento}}')
    ]
    
    for i, (label, placeholder) in enumerate(data):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = placeholder
    
    # Espace
    doc.add_paragraph()
    
    # Section informations supplémentaires
    doc.add_heading('Informações adicionais:', level=1)
    doc.add_paragraph('Data de pedido: _________________')
    doc.add_paragraph('Assinatura: _____________________')
    
    # Espace
    doc.add_paragraph()
    
    # Section observations
    doc.add_heading('Observações:', level=1)
    for i in range(3):
        doc.add_paragraph('_________________________________')
    
    # Espace
    doc.add_paragraph()
    
    # Note de bas de page
    note = doc.add_paragraph('Este documento foi gerado automaticamente pelo sistema eConsulat.')
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

def main():
    """Fonction principale"""
    
    print("🛂 Création d'un template Word simple")
    print("=" * 40)
    
    try:
        # Créer le template
        doc = create_simple_template()
        
        # Chemin de sauvegarde
        template_path = "backend/src/main/resources/FORMULARIO DE PEDIDO DE PASSAPORTE.docx"
        
        # Sauvegarder
        doc.save(template_path)
        
        print(f"✅ Template créé: {template_path}")
        print("📋 Contenu:")
        print("- Titre: FORMULARIO DE PEDIDO DE PASSAPORTE")
        print("- Placeholders: {{Prénom}}, {{Nom de famille}}, {{Date de naissance}}, {{Local de nascimento}}")
        print("- Format: Document Word (.docx)")
        
        return True
        
    except Exception as e:
        print(f"❌ Erreur: {e}")
        return False

if __name__ == "__main__":
    success = main()
    exit(0 if success else 1) 