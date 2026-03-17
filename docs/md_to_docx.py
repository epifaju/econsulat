# -*- coding: utf-8 -*-
"""Convert services-integration-eConsulat.md to DOCX."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_paragraph(doc, text, style=None):
    if not text or not text.strip():
        return
    text = text.strip()
    # Strip markdown bold
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    p = doc.add_paragraph(text, style=style)
    return p

def add_heading(doc, text, level=1):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    doc.add_heading(text.strip(), level=level)

def parse_table(lines):
    """Parse markdown table into list of rows (list of cells). Skip separator line."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line or not line.startswith('|'):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if not cells:
            continue
        # Skip separator row (only dashes and spaces)
        if all(re.match(r'^[\s\-]+$', c) for c in cells):
            continue
        rows.append(cells)
    return rows

def main():
    base = Path(__file__).parent
    md_path = base / "services-integration-eConsulat.md"
    docx_path = base / "services-integration-eConsulat.docx"

    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    doc.add_heading("Services pouvant être intégrés à l'application eConsulat", 0)
    doc.add_paragraph(
        "Document de synthèse des échanges sur les types de services à intégrer, "
        "les courtes listes priorisées, et les architectures détaillées (sans code)."
    )

    lines = content.split("\n")
    i = 0
    table_buffer = []

    while i < len(lines):
        line = lines[i]
        raw = line

        # Skip first title and intro (already added)
        if i < 6 and (line.startswith("# ") or line.startswith("---") or (line.strip() == "")):
            i += 1
            continue

        if line.startswith("# "):
            add_heading(doc, line[2:], level=1)
        elif line.startswith("## "):
            add_heading(doc, line[3:], level=2)
        elif line.startswith("### "):
            add_heading(doc, line[4:], level=3)
        elif line.strip() == "---":
            doc.add_paragraph()
        elif line.strip().startswith("|") and "---" not in line:
            table_buffer = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_buffer.append(lines[i])
                i += 1
            rows = parse_table(table_buffer)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        if ci < len(table.rows[ri].cells):
                            table.rows[ri].cells[ci].text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell)
                doc.add_paragraph()
            table_buffer = []
            continue
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            text = line.strip()[2:].strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            doc.add_paragraph(text, style="List Bullet")
        elif line.strip().startswith("1. ") or (re.match(r"^\d+\.\s", line.strip())):
            text = re.sub(r"^\d+\.\s*", "", line.strip()).strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            doc.add_paragraph(text, style="List Number")
        elif line.strip():
            add_paragraph(doc, line)
        else:
            doc.add_paragraph()

        i += 1

    doc.add_paragraph()
    doc.add_paragraph(
        "Document généré à partir des échanges sur l'intégration de services dans eConsulat. "
        "Aucun code n'a été écrit ; toutes les descriptions sont au niveau architecture et spécification.",
        style="Intense Quote"
    )

    doc.save(docx_path)
    print(f"Created: {docx_path}")

if __name__ == "__main__":
    main()
