#!/usr/bin/env python3
"""
Script pour créer un nouveau template Word valide
Remplace le fichier corrompu FORMULARIO DE PEDIDO DE PASSAPORTE.docx
"""

import os
import sys
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_passport_template():
    """Créer un nouveau template Word pour les passeports"""
    
    # Créer un nouveau document Word
    doc = Document()
    
    # Titre principal
    title = doc.add_heading('FORMULARIO DE PEDIDO DE PASSAPORTE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Espace après le titre
    doc.add_paragraph()
    
    # Section des données personnelles
    doc.add_heading('Dados pessoais:', level=1)
    
    # Tableau pour les données
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    # Remplir le tableau
    data = [
        ('Nome:', '{{Prénom}}'),
        ('Apelido:', '{{Nom de famille}}'),
        ('Data de nascimento:', '{{Date de naissance}}'),
        ('Local de nascimento:', '{{Local de nascimento}}')
    ]
    
    for i, (label, placeholder) in enumerate(data):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = placeholder
    
    # Espace
    doc.add_paragraph()
    
    # Section informations supplémentaires
    doc.add_heading('Informações adicionais:', level=1)
    
    # Date de demande
    doc.add_paragraph('Data de pedido: _________________')
    
    # Signature
    doc.add_paragraph('Assinatura: _____________________')
    
    # Espace
    doc.add_paragraph()
    
    # Section observations
    doc.add_heading('Observações:', level=1)
    for i in range(3):
        doc.add_paragraph('_________________________________')
    
    # Espace
    doc.add_paragraph()
    
    # Note de bas de page
    note = doc.add_paragraph('Este documento foi gerado automaticamente pelo sistema eConsulat.')
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

def main():
    """Fonction principale"""
    
    print("🛂 Création d'un nouveau template Word pour les passeports")
    print("=" * 60)
    
    # Chemin du fichier template
    template_path = "backend/src/main/resources/FORMULARIO DE PEDIDO DE PASSAPORTE.docx"
    
    # Vérifier si le fichier existe
    if os.path.exists(template_path):
        print(f"⚠️  Le fichier existe déjà: {template_path}")
        print("   Il sera remplacé par le nouveau template.")
        
        # Sauvegarder l'ancien fichier
        backup_path = template_path + ".backup"
        try:
            os.rename(template_path, backup_path)
            print(f"✅ Ancien fichier sauvegardé: {backup_path}")
        except Exception as e:
            print(f"❌ Erreur lors de la sauvegarde: {e}")
            return False
    else:
        print(f"📝 Création d'un nouveau fichier: {template_path}")
    
    try:
        # Créer le template
        print("📄 Création du contenu du template...")
        doc = create_passport_template()
        
        # Créer le dossier si nécessaire
        os.makedirs(os.path.dirname(template_path), exist_ok=True)
        
        # Sauvegarder le document
        print(f"💾 Sauvegarde du template...")
        doc.save(template_path)
        
        print(f"✅ Template créé avec succès: {template_path}")
        print()
        print("📋 Contenu du template:")
        print("- Titre: FORMULARIO DE PEDIDO DE PASSAPORTE")
        print("- Placeholders: {{Prénom}}, {{Nom de famille}}, {{Date de naissance}}, {{Local de nascimento}}")
        print("- Format: Document Word (.docx)")
        print()
        print("🔄 Prochaines étapes:")
        print("1. Redémarrez l'application")
        print("2. Testez la génération de document de passeport")
        
        return True
        
    except Exception as e:
        print(f"❌ Erreur lors de la création du template: {e}")
        
        # Restaurer l'ancien fichier si possible
        if os.path.exists(template_path + ".backup"):
            try:
                os.rename(template_path + ".backup", template_path)
                print("🔄 Ancien fichier restauré")
            except:
                pass
        
        return False

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1) 